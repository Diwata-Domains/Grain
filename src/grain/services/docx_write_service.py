"""Writable `.docx` workflow support for Phase 23 office artifacts."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from grain.domain.office_writes import OfficeWriteDecision, OfficeWriteRequest
from grain.services.office_write_service import OfficeWriteService

VALID_DOCX_REPLACEMENT_SCOPES: frozenset[str] = frozenset(
    {"paragraphs", "tables", "paragraphs_and_tables"}
)


@dataclass(frozen=True)
class DocxTextReplacement:
    match_text: str
    replacement_text: str
    scope: str = "paragraphs_and_tables"

    def __post_init__(self) -> None:
        if not self.match_text:
            raise ValueError("DocxTextReplacement.match_text must not be empty")
        if self.scope not in VALID_DOCX_REPLACEMENT_SCOPES:
            raise ValueError(
                f"Invalid docx replacement scope {self.scope!r}. "
                f"Must be one of: {sorted(VALID_DOCX_REPLACEMENT_SCOPES)}"
            )


@dataclass(frozen=True)
class DocxWriteResult:
    decision: OfficeWriteDecision
    output_path: str
    change_summary: list[str] = field(default_factory=list)
    paragraph_replacements: int = 0
    table_cell_replacements: int = 0
    headings_preserved: int = 0
    tables_preserved: int = 0


class DocxWriteService:
    """Apply bounded `.docx` updates via explicit proposal/export outputs."""

    def __init__(self, office_write_service: OfficeWriteService | None = None) -> None:
        self._office_write_service = office_write_service or OfficeWriteService()

    def write_document(
        self,
        *,
        root: Path,
        request: OfficeWriteRequest,
        replacements: list[DocxTextReplacement],
    ) -> DocxWriteResult:
        decision = self._office_write_service.resolve_write_mode(request)
        if decision.resolved_mode == "apply":
            raise ValueError("In-place `.docx` apply mode is not supported in Phase 23")

        source_path = self._resolve(root, request.artifact.source_path)
        output_path = self._resolve_output(root, request.artifact.output_path)
        document = self._load_document(source_path)

        paragraph_replacements = 0
        table_cell_replacements = 0
        headings_preserved = 0

        for paragraph in document.paragraphs:
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if self._heading_level(style_name) > 0:
                headings_preserved += 1
            for replacement in replacements:
                if replacement.scope not in {"paragraphs", "paragraphs_and_tables"}:
                    continue
                updated_text, count = self._replace_text(paragraph.text, replacement)
                if count > 0:
                    paragraph.text = updated_text
                    paragraph_replacements += count

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for replacement in replacements:
                        if replacement.scope not in {"tables", "paragraphs_and_tables"}:
                            continue
                        updated_text, count = self._replace_text(cell.text, replacement)
                        if count > 0:
                            cell.text = updated_text
                            table_cell_replacements += count

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

        change_summary = [
            f"mode: {decision.resolved_mode}",
            f"paragraph replacements: {paragraph_replacements}",
            f"table cell replacements: {table_cell_replacements}",
            f"headings preserved: {headings_preserved}",
            f"tables preserved: {len(document.tables)}",
        ]
        for replacement in replacements:
            change_summary.append(
                f"replacement[{replacement.scope}]: {replacement.match_text!r} -> {replacement.replacement_text!r}"
            )

        return DocxWriteResult(
            decision=decision,
            output_path=str(output_path.relative_to(root)),
            change_summary=change_summary,
            paragraph_replacements=paragraph_replacements,
            table_cell_replacements=table_cell_replacements,
            headings_preserved=headings_preserved,
            tables_preserved=len(document.tables),
        )

    def _resolve(self, root: Path, raw_path: str) -> Path:
        path = Path(raw_path)
        return path if path.is_absolute() else root / path

    def _resolve_output(self, root: Path, raw_path: str) -> Path:
        if not raw_path.strip():
            raise ValueError("`.docx` write outputs require OfficeArtifactRef.output_path")
        return self._resolve(root, raw_path)

    def _load_document(self, path: Path):
        from docx import Document

        return Document(path)

    def _replace_text(
        self,
        original_text: str,
        replacement: DocxTextReplacement,
    ) -> tuple[str, int]:
        count = original_text.count(replacement.match_text)
        if count == 0:
            return original_text, 0
        return (
            original_text.replace(
                replacement.match_text,
                replacement.replacement_text,
            ),
            count,
        )

    def _heading_level(self, style_name: str) -> int:
        lowered = style_name.lower().strip()
        if not lowered.startswith("heading"):
            return 0
        suffix = lowered.replace("heading", "", 1).strip()
        if suffix.isdigit():
            return int(suffix)
        return 1
