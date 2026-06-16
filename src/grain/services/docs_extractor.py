# SPDX-FileCopyrightText: 2024-2026 Shaznay Sison
# SPDX-License-Identifier: AGPL-3.0-only

"""Document extraction service for docs adapter context sources."""

from __future__ import annotations

from pathlib import Path
from typing import Any


class DocsExtractor:
    """Extract readable text from `.docx` and `.md` files."""

    def extract(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".md":
            return self._extract_markdown(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        return f"[docs_extractor: unsupported file type {path.suffix or '(none)'}]"

    def _extract_markdown(self, path: Path) -> str:
        try:
            text = path.read_text(encoding="utf-8")
        except Exception as exc:  # noqa: BLE001 - graceful degradation required
            return f"[docs_extractor: could not read {path.name} - {exc}]"

        if not text.strip():
            return f"[docs_extractor: {path.name} is empty]"
        return text

    def _extract_docx(self, path: Path) -> str:
        try:
            document = self._load_document(path)
        except Exception as exc:  # noqa: BLE001 - graceful degradation required
            return f"[docs_extractor: could not read {path.name} - {exc}]"

        lines: list[str] = [f"# Document: {path.name}"]
        has_content = False

        for paragraph in document.paragraphs:
            text = (paragraph.text or "").strip()
            if not text:
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            heading_level = self._heading_level(style_name)
            if heading_level > 0:
                level = min(6, heading_level + 1)
                lines.append(f"{'#' * level} {text}")
            else:
                lines.append(text)
            has_content = True

        for table_index, table in enumerate(document.tables, start=1):
            if has_content:
                lines.append("")
            lines.append(f"## Table {table_index}")
            for row in table.rows:
                cells = [self._normalize_cell_text(cell.text) for cell in row.cells]
                lines.append(f"| {' | '.join(cells)} |")
                has_content = True

        if not has_content:
            return f"[docs_extractor: {path.name} is empty]"
        return "\n".join(lines)

    def _load_document(self, path: Path) -> Any:
        from docx import Document

        return Document(path)

    def _heading_level(self, style_name: str) -> int:
        lowered = style_name.lower().strip()
        if not lowered.startswith("heading"):
            return 0
        suffix = lowered.replace("heading", "", 1).strip()
        if suffix.isdigit():
            return int(suffix)
        return 1

    def _normalize_cell_text(self, value: str) -> str:
        text = value.replace("\n", " ").strip()
        return text

