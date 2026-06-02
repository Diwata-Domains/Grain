from __future__ import annotations

import json
from pathlib import Path

from click.testing import CliRunner
from docx import Document
from openpyxl import Workbook, load_workbook

from grain.cli import main

_TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "tasks"


def _write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This is a summary paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "owner"
    table.rows[1].cells[1].text = "team-a"
    doc.save(path)


def _create_xlsx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Revenue"
    ws["A1"] = "month"
    ws["B1"] = "amount"
    ws["A2"] = "Jan"
    ws["B2"] = 12
    ws["A3"] = "Feb"
    ws["B3"] = "=SUM(B2,18)"
    wb.save(path)


def _seed_packet_repo(repo: Path) -> None:
    rules = repo / "docs" / "runtime" / "PROJECT_RULES.md"
    rules.parent.mkdir(parents=True, exist_ok=True)
    rules.touch()

    dest_templates = repo / "templates" / "tasks"
    dest_templates.mkdir(parents=True, exist_ok=True)
    for name in ("task.md", "context.md", "plan.md", "deliverable_spec.md", "results.md"):
        src = _TEMPLATES_DIR / name
        (dest_templates / name).write_text(src.read_text(encoding="utf-8"), encoding="utf-8")

    (repo / "tasks").mkdir(parents=True, exist_ok=True)


def _create_packet(repo: Path, phase: int, task_num: int) -> str:
    _seed_packet_repo(repo)
    runner = CliRunner()
    result = runner.invoke(
        main,
        ["--repo", str(repo), "task", "create", "--phase", str(phase), "--task-num", str(task_num)],
    )
    assert result.exit_code == 0, result.output
    packet_dir = next((repo / "tasks").iterdir())
    task_id = packet_dir.name.split("-")[-1]
    (repo / "docs" / "working").mkdir(parents=True, exist_ok=True)
    (repo / "docs" / "working" / "current_task.md").write_text(
        f"# Current Task\n\nTask ID: {task_id}\nTask Path: tasks/{packet_dir.name}/\nStatus: in_progress\n",
        encoding="utf-8",
    )
    return task_id


def test_office_docx_propose_creates_output_and_review_artifact(tmp_path: Path) -> None:
    _create_docx(tmp_path / "docs" / "brief.docx")
    _create_packet(tmp_path, phase=23, task_num=5)

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "docx",
            "propose",
            "--source",
            "docs/brief.docx",
            "--replace",
            "team-a=team-b",
        ],
    )

    assert result.exit_code == 0, result.output
    review_path = next((tmp_path / "tasks").iterdir()) / "office_review.json"
    assert review_path.exists()
    output = next((tmp_path / "tasks").iterdir()) / "brief.proposed.docx"
    assert output.exists()
    output_doc = Document(output)
    assert output_doc.tables[0].rows[1].cells[1].text == "team-b"
    assert "operation_mode   propose" in result.output


def test_office_spreadsheet_export_creates_output_and_review_artifact(tmp_path: Path) -> None:
    _create_xlsx(tmp_path / "data" / "report.xlsx")
    _create_packet(tmp_path, phase=23, task_num=5)

    runner = CliRunner()
    result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "spreadsheet",
            "export",
            "--source",
            "data/report.xlsx",
            "--set",
            "Revenue!B3==SUM(B2,20)",
        ],
    )

    assert result.exit_code == 0, result.output
    output = next((tmp_path / "tasks").iterdir()) / "report.review.xlsx"
    assert output.exists()
    workbook = load_workbook(output, data_only=False)
    assert workbook["Revenue"]["B3"].value == "=SUM(B2,20)"
    review_path = next((tmp_path / "tasks").iterdir()) / "office_review.json"
    payload = json.loads(review_path.read_text(encoding="utf-8"))
    assert payload["review_bundle"]["operation_mode"] == "export-as-new-file"


def test_office_review_show_reports_persisted_review_artifact(tmp_path: Path) -> None:
    _create_docx(tmp_path / "docs" / "brief.docx")
    task_id = _create_packet(tmp_path, phase=23, task_num=5)

    runner = CliRunner()
    invoke_result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "docx",
            "propose",
            "--source",
            "docs/brief.docx",
            "--replace",
            "team-a=team-b",
        ],
    )
    assert invoke_result.exit_code == 0, invoke_result.output

    result = runner.invoke(
        main,
        ["--repo", str(tmp_path), "office", "review", "show", "--task-id", task_id],
    )

    assert result.exit_code == 0, result.output
    assert "office review show: ok" in result.output
    assert "validator_results" in result.output
    assert "docx-structure" in result.output


def test_office_review_show_json_output(tmp_path: Path) -> None:
    _create_xlsx(tmp_path / "data" / "report.xlsx")
    _create_packet(tmp_path, phase=23, task_num=5)

    runner = CliRunner()
    invoke_result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "spreadsheet",
            "propose",
            "--source",
            "data/report.xlsx",
            "--set",
            "Revenue!B2=14",
        ],
    )
    assert invoke_result.exit_code == 0, invoke_result.output

    result = runner.invoke(
        main,
        ["--repo", str(tmp_path), "--format", "json", "office", "review", "show"],
    )

    assert result.exit_code == 0, result.output
    payload = json.loads(result.output)
    assert payload["ok"] is True
    assert payload["office_review"]["review_bundle"]["operation_mode"] == "propose"


def test_office_cli_smoke_flow_uses_active_packet_and_review_artifact(tmp_path: Path) -> None:
    _create_docx(tmp_path / "docs" / "brief.docx")
    _create_xlsx(tmp_path / "data" / "report.xlsx")
    task_id = _create_packet(tmp_path, phase=23, task_num=6)

    runner = CliRunner()

    docx_result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "docx",
            "propose",
            "--source",
            "docs/brief.docx",
            "--replace",
            "team-a=team-b",
        ],
    )
    assert docx_result.exit_code == 0, docx_result.output

    packet_dir = next((tmp_path / "tasks").iterdir())
    proposed_docx = packet_dir / "brief.proposed.docx"
    assert proposed_docx.exists()
    assert Document(proposed_docx).tables[0].rows[1].cells[1].text == "team-b"

    spreadsheet_result = runner.invoke(
        main,
        [
            "--repo",
            str(tmp_path),
            "office",
            "spreadsheet",
            "export",
            "--source",
            "data/report.xlsx",
            "--set",
            "Revenue!B2=14",
        ],
    )
    assert spreadsheet_result.exit_code == 0, spreadsheet_result.output

    exported_xlsx = packet_dir / "report.review.xlsx"
    assert exported_xlsx.exists()
    assert load_workbook(exported_xlsx, data_only=False)["Revenue"]["B2"].value == 14

    review_result = runner.invoke(
        main,
        ["--repo", str(tmp_path), "--format", "json", "office", "review", "show", "--task-id", task_id],
    )
    assert review_result.exit_code == 0, review_result.output

    payload = json.loads(review_result.output)
    assert payload["task_id"] == task_id
    assert payload["office_review"]["review_bundle"]["operation_mode"] == "export-as-new-file"
    assert any(path.endswith("report.review.xlsx") for path in payload["office_review"]["review_bundle"]["artifact_paths"])
    assert any(item["validator_id"] == "spreadsheet-structure" for item in payload["office_review"]["validator_results"])
