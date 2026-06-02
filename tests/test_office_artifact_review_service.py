from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

from grain.domain import OfficeArtifactRef, OfficeWriteRequest
from grain.services.docx_write_service import DocxTextReplacement, DocxWriteService
from grain.services.office_artifact_review_service import OfficeArtifactReviewService
from grain.services.spreadsheet_write_service import (
    SpreadsheetCellUpdate,
    SpreadsheetWriteService,
)


def _create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This is a summary paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "owner"
    table.rows[1].cells[1].text = "team-a"
    doc.save(path)


def _create_xlsx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Revenue"
    ws["A1"] = "month"
    ws["B1"] = "amount"
    ws["A2"] = "Jan"
    ws["B2"] = 12
    ws["A3"] = "Feb"
    ws["B3"] = "=SUM(B2,18)"
    wb.save(path)


def test_build_docx_review_result_collects_all_validator_categories(
    tmp_path: Path,
) -> None:
    source = tmp_path / "docs" / "brief.docx"
    _create_docx(source)
    write_result = DocxWriteService().write_document(
        root=tmp_path,
        request=OfficeWriteRequest(
            packet_id="TASK-0155",
            artifact=OfficeArtifactRef(
                "docx",
                "docs/brief.docx",
                output_path="tasks/P23-T04-TASK-0155/brief.proposed.docx",
            ),
            requested_mode="propose",
        ),
        replacements=[DocxTextReplacement("team-a", "team-b", scope="tables")],
    )

    review_result = OfficeArtifactReviewService().build_docx_review_result(
        root=tmp_path,
        packet_id="TASK-0155",
        result=write_result,
    )

    assert [item.category for item in review_result.validator_results] == [
        "structure",
        "reference",
        "policy",
    ]
    assert all(item.state == "passed" for item in review_result.validator_results)
    assert review_result.review_bundle.operation_mode == "propose"
    assert review_result.review_bundle.artifact_paths == [
        "docs/brief.docx",
        "tasks/P23-T04-TASK-0155/brief.proposed.docx",
    ]
    assert review_result.review_bundle.residual_risks == []


def test_build_spreadsheet_review_result_collects_formula_aware_summary(
    tmp_path: Path,
) -> None:
    source = tmp_path / "data" / "report.xlsx"
    _create_xlsx(source)
    write_result = SpreadsheetWriteService().write_workbook(
        root=tmp_path,
        request=OfficeWriteRequest(
            packet_id="TASK-0155",
            artifact=OfficeArtifactRef(
                "spreadsheet",
                "data/report.xlsx",
                output_path="tasks/P23-T04-TASK-0155/report.proposed.xlsx",
            ),
            requested_mode="export-as-new-file",
        ),
        updates=[SpreadsheetCellUpdate("Revenue", "B3", "=SUM(B2,20)")],
    )

    review_result = OfficeArtifactReviewService().build_spreadsheet_review_result(
        root=tmp_path,
        packet_id="TASK-0155",
        result=write_result,
    )

    assert [item.category for item in review_result.validator_results] == [
        "structure",
        "reference",
        "policy",
    ]
    assert all(item.state == "passed" for item in review_result.validator_results)
    assert review_result.review_bundle.operation_mode == "export-as-new-file"
    assert "formula cells changed: Revenue!B3" in review_result.review_bundle.change_summary


def test_missing_output_becomes_reference_failure_and_residual_risk(
    tmp_path: Path,
) -> None:
    source = tmp_path / "docs" / "brief.docx"
    _create_docx(source)
    write_result = DocxWriteService().write_document(
        root=tmp_path,
        request=OfficeWriteRequest(
            packet_id="TASK-0155",
            artifact=OfficeArtifactRef(
                "docx",
                "docs/brief.docx",
                output_path="tasks/P23-T04-TASK-0155/brief.proposed.docx",
            ),
            requested_mode="propose",
        ),
        replacements=[DocxTextReplacement("Scope", "Updated Scope")],
    )
    (tmp_path / write_result.output_path).unlink()

    review_result = OfficeArtifactReviewService().build_docx_review_result(
        root=tmp_path,
        packet_id="TASK-0155",
        result=write_result,
    )

    reference_validator = next(
        item for item in review_result.validator_results if item.category == "reference"
    )
    assert reference_validator.state == "failed"
    assert "missing output file" in reference_validator.findings[0]
    assert "one or more validators failed" in review_result.review_bundle.residual_risks
