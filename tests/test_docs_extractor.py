"""Tests for docs extraction and docs adapter context wiring."""

from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document

from grain.adapters.export import render_context_markdown_export
from grain.services.context_service import build_context_bundle
from grain.services.docs_extractor import DocsExtractor
from grain.services.task_service import create_packet_directory


def _write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _write_manifest(root: Path) -> None:
    manifest = {
        "canonical": [
            {
                "id": "workflow_spec",
                "path": "docs/canonical/workflow_spec.md",
                "purpose": "workflow",
                "authority": "highest",
                "editable_by_agents": False,
                "read_when": ["running_tasks"],
            }
        ],
        "working": [],
        "runtime": [],
        "tasks": {},
        "rules": {},
    }
    path = root / "docs" / "runtime" / "docs_manifest.yaml"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.dump(manifest), encoding="utf-8")
    _write(root / "docs" / "canonical" / "workflow_spec.md", "# Workflow Spec\n")


def _write_docs_adapter_profile(root: Path) -> None:
    _write(
        root / "docs" / "runtime" / "adapter_profiles.md",
        """# Adapter Profiles

## 5. Adapter Profiles

### docs_adapter
- `adapter_id`: `docs_adapter`
- `domain_type`: `docs`
- `applies_to`:
  - markdown
  - docx
- `relevant_file_patterns`:
  - `**/*.md`
  - `**/*.docx`
- `context_priority_rules`:
  - prioritize canonical docs first
- `test_or_validation_hints`:
  - validate requirement headings
""",
    )


def _set_primary_adapter(root: Path, packet_dir_name: str, adapter_id: str) -> None:
    task_md = root / "tasks" / packet_dir_name / "task.md"
    task_md.write_text(
        task_md.read_text(encoding="utf-8").replace(
            "- **Primary Adapter:** none",
            f"- **Primary Adapter:** {adapter_id}",
        ),
        encoding="utf-8",
    )


def _create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This is a summary paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "owner"
    table.rows[1].cells[1].text = "team-a"
    doc.save(path)


def test_extract_docx_includes_headings_paragraphs_and_tables(tmp_path: Path):
    path = tmp_path / "docs" / "brief.docx"
    _create_docx(path)

    text = DocsExtractor().extract(path)

    assert "# Document: brief.docx" in text
    assert "## Scope" in text
    assert "This is a summary paragraph." in text
    assert "| Key | Value |" in text
    assert "| owner | team-a |" in text


def test_extract_markdown_returns_raw_text(tmp_path: Path):
    path = tmp_path / "docs" / "notes.md"
    _write(path, "# Notes\nhello\n")

    text = DocsExtractor().extract(path)

    assert text == "# Notes\nhello\n"


def test_extract_empty_markdown_returns_empty_marker(tmp_path: Path):
    path = tmp_path / "docs" / "empty.md"
    _write(path, "")

    text = DocsExtractor().extract(path)
    assert "empty.md is empty" in text


def test_extract_unreadable_docx_returns_warning(tmp_path: Path):
    text = DocsExtractor().extract(tmp_path / "missing.docx")
    assert "could not read missing.docx" in text


def test_extract_empty_docx_returns_empty_marker(tmp_path: Path):
    path = tmp_path / "docs" / "empty.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.save(path)

    text = DocsExtractor().extract(path)
    assert "empty.docx is empty" in text


def test_context_bundle_selects_docx_and_md_sources_with_docs_adapter(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_docs_adapter_profile(packet_repo)
    _create_docx(packet_repo / "docs" / "brief.docx")
    _write(packet_repo / "docs" / "notes.md", "# Notes\n")
    create_packet_directory(packet_repo, phase=14, task_num=2)
    _set_primary_adapter(packet_repo, "P14-T02-TASK-0001", "docs_adapter")

    result, bundle = build_context_bundle(packet_repo, "TASK-0001")
    assert result.ok is True
    assert bundle is not None
    sources = bundle.export_metadata["sources"]
    assert "docs/brief.docx" in sources
    assert "docs/notes.md" in sources


def test_context_export_renders_docx_extracted_content(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_docs_adapter_profile(packet_repo)
    _create_docx(packet_repo / "docs" / "brief.docx")
    create_packet_directory(packet_repo, phase=14, task_num=2)
    _set_primary_adapter(packet_repo, "P14-T02-TASK-0001", "docs_adapter")

    result, bundle = build_context_bundle(packet_repo, "TASK-0001")
    assert result.ok is True
    assert bundle is not None

    content = render_context_markdown_export(packet_repo, bundle)
    assert "## Source: `docs/brief.docx`" in content
    assert "## Scope" in content
    assert "| Key | Value |" in content


def test_context_export_keeps_markdown_source_rendering(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_docs_adapter_profile(packet_repo)
    _write(packet_repo / "docs" / "notes.md", "# Notes\nline\n")
    create_packet_directory(packet_repo, phase=14, task_num=2)
    _set_primary_adapter(packet_repo, "P14-T02-TASK-0001", "docs_adapter")

    result, bundle = build_context_bundle(packet_repo, "TASK-0001")
    assert result.ok is True
    assert bundle is not None

    content = render_context_markdown_export(packet_repo, bundle)
    assert "## Source: `docs/notes.md`" in content
    assert "# Notes" in content
    assert "line" in content

