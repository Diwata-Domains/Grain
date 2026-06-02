"""Cross-adapter integration tests for Phase 14 document extraction."""

from __future__ import annotations

import csv
import json
from pathlib import Path

import yaml
from click.testing import CliRunner
from docx import Document
from openpyxl import Workbook

from grain.cli import main
from grain.services.task_service import create_packet_directory


def _run(repo_root: Path, *args: str):
    runner = CliRunner()
    return runner.invoke(main, ["--repo", str(repo_root), *args])


def _write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def _write_manifest(root: Path) -> None:
    manifest = {
        "canonical": [
            {
                "id": "workflow_spec",
                "path": "docs/canonical/workflow_spec.md",
                "purpose": "workflow",
                "authority": "highest",
                "editable_by_agents": False,
                "read_when": ["running_tasks"],
            }
        ],
        "working": [],
        "runtime": [],
        "tasks": {},
        "rules": {},
    }
    path = root / "docs" / "runtime" / "docs_manifest.yaml"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.dump(manifest), encoding="utf-8")
    _write(root / "docs" / "canonical" / "workflow_spec.md", "# Workflow Spec\n")


def _write_adapter_profiles(root: Path) -> None:
    _write(
        root / "docs" / "runtime" / "adapter_profiles.md",
        """# Adapter Profiles

## 5. Adapter Profiles

### code_adapter
- `adapter_id`: `code_adapter`
- `domain_type`: `code`
- `applies_to`:
  - Python
- `relevant_file_patterns`:
  - `src/**`
- `test_or_validation_hints`:
  - run focused tests

### spreadsheet_adapter
- `adapter_id`: `spreadsheet_adapter`
- `domain_type`: `data`
- `applies_to`:
  - spreadsheets
- `relevant_file_patterns`:
  - `**/*.xlsx`
  - `**/*.csv`
- `context_priority_rules`:
  - prioritize sheets with explicit headers
- `test_or_validation_hints`:
  - confirm headers and row counts

### docs_adapter
- `adapter_id`: `docs_adapter`
- `domain_type`: `docs`
- `applies_to`:
  - markdown
  - docx
  - pdf
- `relevant_file_patterns`:
  - `**/*.md`
  - `**/*.docx`
  - `**/*.pdf`
- `context_priority_rules`:
  - prioritize canonical docs first
- `test_or_validation_hints`:
  - validate key headings

### obsidian_adapter
- `adapter_id`: `obsidian_adapter`
- `domain_type`: `docs`
- `applies_to`:
  - Obsidian vaults
  - wiki-link-driven markdown
- `relevant_file_patterns`:
  - `**/*.md`
- `ignore_file_patterns`:
  - `tasks/**`
  - `templates/**`
  - `docs/**`
- `context_priority_rules`:
  - prioritize note directories and linked notes before unrelated markdown
- `test_or_validation_hints`:
  - verify wiki-link references and frontmatter presence

### database_adapter
- `adapter_id`: `database_adapter`
- `domain_type`: `code`
- `applies_to`:
  - relational database workflows
  - schema and migration planning
- `relevant_file_patterns`:
  - `**/*.sql`
  - `migrations/**`
  - `queries/**`
  - `models/**`
  - `src/models/**`
  - `repositories/**`
  - `src/**/models/**`
  - `src/db/**`
  - `src/repositories/**`
  - `src/**/db/**`
- `ignore_file_patterns`:
  - `tasks/**`
  - `templates/**`
  - `docs/**`
- `context_priority_rules`:
  - prioritize schema files and migration directories before unrelated application code
  - treat query files, ORM models, and repository layers as secondary context when the task points at persistence behavior
- `test_or_validation_hints`:
  - verify migration ordering and schema consistency

### crawler_adapter
- `adapter_id`: `crawler_adapter`
- `domain_type`: `code`
- `applies_to`:
  - crawler workflows
  - selector and extraction planning
- `relevant_file_patterns`:
  - `crawler/**`
  - `selectors/**`
  - `schemas/**`
  - `fixtures/**`
  - `outputs/**`
  - `normalizers/**`
  - `**/*crawl*.yml`
  - `**/*selector*.yml`
  - `**/*schema*.json`
- `ignore_file_patterns`:
  - `tasks/**`
  - `templates/**`
  - `docs/**`
- `context_priority_rules`:
  - prioritize crawl configs, selectors, and extraction schemas before unrelated application code
  - treat output fixtures and normalization logic as secondary context when the task is about extraction quality
- `test_or_validation_hints`:
  - verify selector coverage and extraction-schema consistency
""",
    )


def _set_primary_adapter(root: Path, packet_dir_name: str, adapter_id: str) -> None:
    task_md = root / "tasks" / packet_dir_name / "task.md"
    task_md.write_text(
        task_md.read_text(encoding="utf-8").replace(
            "- **Primary Adapter:** none",
            f"- **Primary Adapter:** {adapter_id}",
        ),
        encoding="utf-8",
    )


def _set_objective(root: Path, packet_dir_name: str, objective: str) -> None:
    task_md = root / "tasks" / packet_dir_name / "task.md"
    task_md.write_text(
        task_md.read_text(encoding="utf-8").replace(
            "[What must be built or done. One clear paragraph.]",
            objective,
        ),
        encoding="utf-8",
    )


def _create_xlsx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["k", "v"])
    ws.append(["a", 1])
    ws.append(["b", 2])
    wb.save(path)


def _create_csv(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["name", "score"])
        writer.writerow(["alice", "10"])
        writer.writerow(["bob", "12"])


def _create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Spec", level=1)
    doc.add_paragraph("Document paragraph text.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "k"
    table.rows[0].cells[1].text = "v"
    table.rows[1].cells[0].text = "owner"
    table.rows[1].cells[1].text = "team"
    doc.save(path)


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _build_pdf(pages: list[str | None]) -> bytes:
    objects: list[str] = []
    kids: list[str] = []
    objects.append("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    objects.append("2 0 obj\n<< /Type /Pages /Count {count} /Kids [{kids}] >>\nendobj\n")
    objects.append("3 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n")

    obj_id = 4
    for page_text in pages:
        page_obj = obj_id
        content_obj = obj_id + 1
        kids.append(f"{page_obj} 0 R")
        objects.append(
            f"{page_obj} 0 obj\n"
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_obj} 0 R >>\n"
            "endobj\n"
        )
        if page_text is None:
            stream_text = ""
        else:
            stream_text = f"BT /F1 12 Tf 72 720 Td ({_escape_pdf_text(page_text)}) Tj ET"
        objects.append(
            f"{content_obj} 0 obj\n"
            f"<< /Length {len(stream_text.encode('latin-1'))} >>\n"
            "stream\n"
            f"{stream_text}\n"
            "endstream\n"
            "endobj\n"
        )
        obj_id += 2

    objects[1] = objects[1].format(count=len(pages), kids=" ".join(kids))
    data = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    offsets = [0]
    for obj in objects:
        offsets.append(len(data))
        data += obj.encode("latin-1")
    xref_pos = len(data)
    data += f"xref\n0 {len(offsets)}\n".encode("latin-1")
    data += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        data += f"{off:010d} 00000 n \n".encode("latin-1")
    data += (
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF\n"
    ).encode("latin-1")
    return data


def _create_pdf(path: Path, pages: list[str | None]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(_build_pdf(pages))


def _setup_packet(root: Path, *, phase: int, task_num: int, primary_adapter: str) -> None:
    create_packet_directory(root, phase=phase, task_num=task_num)
    packet_name = f"P{phase}-T{task_num:02d}-TASK-0001"
    _set_primary_adapter(root, packet_name, primary_adapter)


def test_context_build_spreadsheet_adapter_selects_xlsx_and_csv(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_xlsx(packet_repo / "data" / "sheet.xlsx")
    _create_csv(packet_repo / "data" / "scores.csv")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="spreadsheet_adapter")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]

    assert result.exit_code == 0, result.output
    assert "data/sheet.xlsx" in sources
    assert "data/scores.csv" in sources


def test_context_export_spreadsheet_adapter_renders_xlsx_content(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_xlsx(packet_repo / "data" / "sheet.xlsx")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="spreadsheet_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `data/sheet.xlsx`" in export
    assert "## Sheet: Data" in export


def test_context_export_spreadsheet_adapter_renders_csv_content(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_csv(packet_repo / "data" / "scores.csv")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="spreadsheet_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `data/scores.csv`" in export
    assert "Columns: name, score" in export


def test_context_build_docs_adapter_selects_docx_md_pdf(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_docx(packet_repo / "docs" / "brief.docx")
    _write(packet_repo / "docs" / "notes.md", "# Notes\n")
    _create_pdf(packet_repo / "docs" / "spec.pdf", ["PDF content"])
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]

    assert result.exit_code == 0, result.output
    assert "docs/brief.docx" in sources
    assert "docs/notes.md" in sources
    assert "docs/spec.pdf" in sources


def test_context_export_docs_adapter_renders_docx_content(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_docx(packet_repo / "docs" / "brief.docx")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `docs/brief.docx`" in export
    assert "## Spec" in export
    assert "| k | v |" in export


def test_context_export_docs_adapter_renders_pdf_content(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_pdf(packet_repo / "docs" / "spec.pdf", ["Text from pdf"])
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `docs/spec.pdf`" in export
    assert "Text from pdf" in export


def test_context_export_corrupt_pdf_degrades_without_failing_bundle(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    bad = packet_repo / "docs" / "broken.pdf"
    bad.parent.mkdir(parents=True, exist_ok=True)
    bad.write_bytes(b"not-pdf")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `docs/broken.pdf`" in export
    assert "pdf_extractor: could not read broken.pdf" in export


def test_context_build_text_output_succeeds_for_docs_adapter(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_docx(packet_repo / "docs" / "brief.docx")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "context", "build", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    assert "context build: ok" in result.output
    assert "primary_adapter   docs_adapter" in result.output


def test_context_build_mixed_repo_code_and_docs_bundle_succeeds(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "src" / "main.py", "print('ok')\n")
    _create_docx(packet_repo / "docs" / "brief.docx")
    _create_pdf(packet_repo / "docs" / "spec.pdf", ["doc page"])
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    stats = payload["context_stats"]
    assert result.exit_code == 0, result.output
    assert stats["total_sources"] >= 7
    assert "docs/brief.docx" in payload["bundle"]["export_metadata"]["sources"]
    assert "docs/spec.pdf" in payload["bundle"]["export_metadata"]["sources"]


def test_context_stats_counts_spreadsheet_document_sources(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _create_xlsx(packet_repo / "data" / "sheet.xlsx")
    _create_csv(packet_repo / "data" / "scores.csv")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="spreadsheet_adapter")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    stats = payload["context_stats"]
    document_entries = [
        item
        for item in stats["per_file"]
        if item["path"] in {"data/sheet.xlsx", "data/scores.csv"}
    ]
    assert result.exit_code == 0, result.output
    assert len(document_entries) == 2
    assert all(item["selection_method"] == "glob_only" for item in document_entries)


def test_context_stats_counts_docs_adapter_document_sources(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "docs" / "notes.md", "# Notes\n")
    _create_docx(packet_repo / "docs" / "brief.docx")
    _create_pdf(packet_repo / "docs" / "spec.pdf", ["pdf text"])
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    stats = payload["context_stats"]
    document_entries = [
        item
        for item in stats["per_file"]
        if item["path"] in {"docs/notes.md", "docs/brief.docx", "docs/spec.pdf"}
    ]
    assert result.exit_code == 0, result.output
    assert len(document_entries) == 3
    assert all(item["selection_method"] == "glob_only" for item in document_entries)


def test_context_export_all_three_document_types_in_one_bundle(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "docs" / "notes.md", "# Notes\n")
    _create_docx(packet_repo / "docs" / "brief.docx")
    _create_pdf(packet_repo / "docs" / "spec.pdf", ["all-types"])
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="docs_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `docs/notes.md`" in export
    assert "Source: `docs/brief.docx`" in export
    assert "Source: `docs/spec.pdf`" in export


def test_context_build_obsidian_adapter_selects_vault_markdown(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(
        packet_repo / "vault" / "Projects" / "Launch Plan.md",
        "---\ntags: [launch]\n---\n# Launch Plan\nSee [[Daily Note]].\n",
    )
    _write(packet_repo / "vault" / "Daily" / "Daily Note.md", "# Daily Note\nShip checklist.\n")
    _write(packet_repo / "vault" / "Archive" / "Archive Note.md", "# Archive Note\nOld note.\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="obsidian_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Update the Launch Plan note for the vault.")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]
    per_file = payload["context_stats"]["per_file"]

    assert result.exit_code == 0, result.output
    assert "vault/Projects/Launch Plan.md" in sources
    assert "vault/Daily/Daily Note.md" in sources
    assert sources.index("vault/Projects/Launch Plan.md") < sources.index("vault/Daily/Daily Note.md")
    assert sources.index("vault/Daily/Daily Note.md") < sources.index("vault/Archive/Archive Note.md")
    linked_entry = next(item for item in per_file if item["path"] == "vault/Daily/Daily Note.md")
    assert linked_entry["selection_method"] == "wiki_linked"


def test_context_export_obsidian_adapter_renders_frontmatter_and_wikilinks(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(
        packet_repo / "vault" / "Projects" / "Launch Plan.md",
        "---\ntags: [launch]\nowner: team-a\n---\n# Launch Plan\nSee [[Daily Note]].\n",
    )
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="obsidian_adapter")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")
    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `vault/Projects/Launch Plan.md`" in export
    assert "tags: [launch]" in export
    assert "[[Daily Note]]" in export


def test_context_build_database_adapter_prioritizes_schema_migrations_and_models(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "database" / "schema.sql", "create table users(id integer primary key);\n")
    _write(packet_repo / "migrations" / "001_create_users.sql", "-- create users table\n")
    _write(packet_repo / "src" / "models" / "user.py", "class User:\n    pass\n")
    _write(packet_repo / "src" / "main.py", "print('app')\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="database_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Plan the schema migration for the users table.")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]
    per_file = payload["context_stats"]["per_file"]

    assert result.exit_code == 0, result.output
    assert "database/schema.sql" in sources
    assert "migrations/001_create_users.sql" in sources
    assert "src/models/user.py" in sources
    assert "src/main.py" not in sources
    assert sources.index("database/schema.sql") < sources.index("src/models/user.py")
    assert sources.index("migrations/001_create_users.sql") < sources.index("src/models/user.py")
    model_entry = next(item for item in per_file if item["path"] == "src/models/user.py")
    assert model_entry["selection_method"] == "glob_only"


def test_context_build_database_adapter_includes_query_and_repository_surfaces_for_persistence_objective(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "database" / "schema.sql", "create table users(id integer primary key);\n")
    _write(packet_repo / "queries" / "list_users.sql", "select * from users;\n")
    _write(packet_repo / "src" / "repositories" / "user_repo.py", "def list_users():\n    return []\n")
    _write(packet_repo / "src" / "models" / "user.py", "class User:\n    pass\n")
    _write(packet_repo / "src" / "main.py", "print('app')\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="database_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Review the query and repository persistence flow for listing users.")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]

    assert result.exit_code == 0, result.output
    assert "queries/list_users.sql" in sources
    assert "src/repositories/user_repo.py" in sources
    assert "src/main.py" not in sources
    assert sources.index("queries/list_users.sql") < sources.index("src/models/user.py")
    assert sources.index("src/repositories/user_repo.py") < sources.index("src/models/user.py")


def test_context_export_database_adapter_smoke_flow(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "database" / "schema.sql", "create table users(id integer primary key);\n")
    _write(packet_repo / "migrations" / "001_create_users.sql", "-- create users table\n")
    _write(packet_repo / "queries" / "list_users.sql", "select * from users;\n")
    _write(packet_repo / "src" / "repositories" / "user_repo.py", "def list_users():\n    return []\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="database_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Review the query and migration flow for listing users.")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")

    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `database/schema.sql`" in export
    assert "Source: `migrations/001_create_users.sql`" in export
    assert "Source: `queries/list_users.sql`" in export
    assert "Source: `src/repositories/user_repo.py`" in export


def test_context_build_crawler_adapter_prioritizes_configs_selectors_and_schemas(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "crawler" / "crawl_plan.yml", "start_urls:\n  - https://example.com\n")
    _write(packet_repo / "selectors" / "article_selector.yml", "title: h1\n")
    _write(packet_repo / "schemas" / "article.schema.json", "{\"title\": \"string\"}\n")
    _write(packet_repo / "src" / "main.py", "print('app')\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="crawler_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Review the crawl config and selector plan for the article extractor.")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]
    per_file = payload["context_stats"]["per_file"]

    assert result.exit_code == 0, result.output
    assert "crawler/crawl_plan.yml" in sources
    assert "selectors/article_selector.yml" in sources
    assert "schemas/article.schema.json" in sources
    assert "src/main.py" not in sources
    assert sources.index("crawler/crawl_plan.yml") < sources.index("schemas/article.schema.json")
    assert sources.index("selectors/article_selector.yml") < sources.index("schemas/article.schema.json")
    schema_entry = next(item for item in per_file if item["path"] == "schemas/article.schema.json")
    assert schema_entry["selection_method"] == "glob_only"


def test_context_build_crawler_adapter_includes_output_validation_and_normalization_for_quality_objective(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "crawler" / "crawl_plan.yml", "start_urls:\n  - https://example.com\n")
    _write(packet_repo / "selectors" / "article_selector.yml", "title: h1\n")
    _write(packet_repo / "schemas" / "article.schema.json", "{\"title\": \"string\"}\n")
    _write(packet_repo / "outputs" / "article_fixture.json", "{\"title\": \"Example\"}\n")
    _write(packet_repo / "normalizers" / "article_normalizer.py", "def normalize(item):\n    return item\n")
    _write(packet_repo / "src" / "main.py", "print('app')\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="crawler_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Validate extraction quality and normalize the article output schema.")

    result = _run(packet_repo, "--format", "json", "context", "build", "--id", "TASK-0001")
    payload = json.loads(result.output)
    sources = payload["bundle"]["export_metadata"]["sources"]

    assert result.exit_code == 0, result.output
    assert "outputs/article_fixture.json" in sources
    assert "normalizers/article_normalizer.py" in sources
    assert "src/main.py" not in sources
    assert sources.index("outputs/article_fixture.json") < sources.index("schemas/article.schema.json")
    assert sources.index("normalizers/article_normalizer.py") < sources.index("schemas/article.schema.json")


def test_context_export_crawler_adapter_smoke_flow(packet_repo: Path):
    _write_manifest(packet_repo)
    _write_adapter_profiles(packet_repo)
    _write(packet_repo / "crawler" / "crawl_plan.yml", "start_urls:\n  - https://example.com\n")
    _write(packet_repo / "selectors" / "article_selector.yml", "title: h1\n")
    _write(packet_repo / "schemas" / "article.schema.json", "{\"title\": \"string\"}\n")
    _write(packet_repo / "outputs" / "article_fixture.json", "{\"title\": \"Example\"}\n")
    _write(packet_repo / "normalizers" / "article_normalizer.py", "def normalize(item):\n    return item\n")
    _setup_packet(packet_repo, phase=14, task_num=4, primary_adapter="crawler_adapter")
    _set_objective(packet_repo, "P14-T04-TASK-0001", "Validate extraction quality and review the crawl plan for article outputs.")

    result = _run(packet_repo, "context", "export", "--id", "TASK-0001")

    assert result.exit_code == 0, result.output
    export = (packet_repo / "tasks" / "P14-T04-TASK-0001" / "context_export.md").read_text(
        encoding="utf-8"
    )
    assert "Source: `crawler/crawl_plan.yml`" in export
    assert "Source: `selectors/article_selector.yml`" in export
    assert "Source: `schemas/article.schema.json`" in export
    assert "Source: `outputs/article_fixture.json`" in export
    assert "Source: `normalizers/article_normalizer.py`" in export
