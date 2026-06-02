from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from grain.domain import OfficeArtifactRef, OfficeWriteRequest
from grain.services.docx_write_service import DocxTextReplacement, DocxWriteService


def _create_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This is a summary paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "owner"
    table.rows[1].cells[1].text = "team-a"
    doc.save(path)


def test_docx_propose_writes_reviewable_output_without_mutating_source(tmp_path: Path) -> None:
    source = tmp_path / "docs" / "brief.docx"
    output = tmp_path / "tasks" / "P23-T02-TASK-0153" / "brief.proposed.docx"
    _create_docx(source)

    service = DocxWriteService()
    request = OfficeWriteRequest(
        packet_id="TASK-0153",
        artifact=OfficeArtifactRef(
            "docx",
            "docs/brief.docx",
            output_path="tasks/P23-T02-TASK-0153/brief.proposed.docx",
        ),
        requested_mode="propose",
    )

    result = service.write_document(
        root=tmp_path,
        request=request,
        replacements=[
            DocxTextReplacement("summary paragraph", "updated summary paragraph"),
            DocxTextReplacement("team-a", "team-b", scope="tables"),
        ],
    )

    source_doc = Document(source)
    output_doc = Document(output)

    assert "summary paragraph" in source_doc.paragraphs[1].text
    assert "updated summary paragraph" in output_doc.paragraphs[1].text
    assert output_doc.tables[0].rows[1].cells[1].text == "team-b"
    assert result.output_path == "tasks/P23-T02-TASK-0153/brief.proposed.docx"
    assert result.paragraph_replacements == 1
    assert result.table_cell_replacements == 1
    assert result.headings_preserved == 1
    assert result.tables_preserved == 1
    assert "mode: propose" in result.change_summary


def test_docx_export_mode_persists_export_file_and_summary(tmp_path: Path) -> None:
    source = tmp_path / "docs" / "brief.docx"
    output = tmp_path / "exports" / "brief.review.docx"
    _create_docx(source)

    service = DocxWriteService()
    request = OfficeWriteRequest(
        packet_id="TASK-0153",
        artifact=OfficeArtifactRef(
            "docx",
            "docs/brief.docx",
            output_path="exports/brief.review.docx",
        ),
        requested_mode="export-as-new-file",
    )

    result = service.write_document(
        root=tmp_path,
        request=request,
        replacements=[DocxTextReplacement("Scope", "Updated Scope", scope="paragraphs")],
    )

    output_doc = Document(output)

    assert output.exists()
    assert output_doc.paragraphs[0].text == "Updated Scope"
    assert "mode: export-as-new-file" in result.change_summary
    assert "replacement[paragraphs]: 'Scope' -> 'Updated Scope'" in result.change_summary


def test_docx_apply_mode_is_rejected_for_phase_23(tmp_path: Path) -> None:
    source = tmp_path / "docs" / "brief.docx"
    _create_docx(source)

    service = DocxWriteService()
    request = OfficeWriteRequest(
        packet_id="TASK-0153",
        artifact=OfficeArtifactRef(
            "docx",
            "docs/brief.docx",
            output_path="docs/brief.updated.docx",
        ),
        requested_mode="apply",
        explicit_apply=True,
        validation_state="passed",
    )

    with pytest.raises(ValueError, match="In-place `.docx` apply mode is not supported"):
        service.write_document(
            root=tmp_path,
            request=request,
            replacements=[DocxTextReplacement("Scope", "Updated Scope")],
        )


def test_docx_requires_explicit_output_path(tmp_path: Path) -> None:
    source = tmp_path / "docs" / "brief.docx"
    _create_docx(source)

    service = DocxWriteService()
    request = OfficeWriteRequest(
        packet_id="TASK-0153",
        artifact=OfficeArtifactRef("docx", "docs/brief.docx"),
        requested_mode="propose",
    )

    with pytest.raises(ValueError, match="outputs require OfficeArtifactRef.output_path"):
        service.write_document(
            root=tmp_path,
            request=request,
            replacements=[DocxTextReplacement("Scope", "Updated Scope")],
        )


def test_invalid_docx_replacement_scope_raises() -> None:
    with pytest.raises(ValueError, match="Invalid docx replacement scope"):
        DocxTextReplacement("a", "b", scope="headers")
